"""
Template Processor Service

Procesa plantillas Word (.docx) reemplazando marcadores con datos del evento.
"""

from docx import Document
from typing import Dict
import os


class TemplateProcessor:
    """
    Procesa plantillas Word reemplazando campos editables con datos del evento.
    
    Marcadores soportados:
    - {{fecha_evento}}: Fecha del evento
    - {{hora_evento}}: Hora del evento
    - {{lugar_evento}}: Lugar del evento
    - {{referencia_evento}}: Referencia opcional del evento
    """
    
    def __init__(self, template_dir: str = None):
        """
        Inicializa el procesador con la configuración de plantillas.
        
        Args:
            template_dir: Directorio donde se encuentran las plantillas.
                         Si es None, usa backend/templates/.
        """
        if template_dir is None:
            # Por defecto, las plantillas están en backend/templates/
            current_dir = os.path.dirname(os.path.abspath(__file__))
            backend_dir = os.path.dirname(current_dir)
            self.template_dir = os.path.join(backend_dir, "templates")
        else:
            self.template_dir = template_dir
        
        self.templates = {
            "a4": "Formato a4.docx",
            "4x1": "Formato 4x1.docx"
        }
    
    def process_template(self, template_type: str, data: Dict[str, str], output_path: str) -> str:
        """
        Procesa una plantilla Word reemplazando campos con datos.
        
        Args:
            template_type: Tipo de plantilla ("a4" o "4x1")
            data: Diccionario con los datos del evento:
                - fecha_evento: Fecha del evento (requerido)
                - hora_evento: Hora del evento (requerido)
                - lugar_evento: Lugar del evento (requerido)
                - referencia_evento: Referencia opcional (puede ser vacío)
            output_path: Ruta donde guardar el documento procesado
        
        Returns:
            Ruta del archivo procesado
            
        Raises:
            FileNotFoundError: Si la plantilla no existe
            KeyError: Si template_type no es válido
        """
        if template_type not in self.templates:
            raise KeyError(f"Tipo de plantilla inválido: {template_type}")
        
        template_path = os.path.join(self.template_dir, self.templates[template_type])
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")
        
        # Cargar el documento
        doc = Document(template_path)
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            self._replace_in_runs(paragraph.runs, data)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_runs(paragraph.runs, data)
        
        # Reemplazar en cuadros de texto (textboxes)
        self._replace_in_textboxes(doc, data)
        
        # Guardar el documento procesado
        doc.save(output_path)
        
        return output_path
    
    def _replace_in_runs(self, runs, data: Dict[str, str]):
        """
        Reemplaza marcadores en runs de texto.
        
        Los marcadores tienen el formato {{nombre_campo}} y son reemplazados
        con los valores correspondientes del diccionario data.
        
        Word puede dividir los marcadores en múltiples runs. Esta función
        detecta marcadores fragmentados y los reemplaza sin afectar runs
        que contienen imágenes, saltos de línea u otros elementos.
        
        Args:
            runs: Lista de runs de texto de un párrafo
            data: Diccionario con los valores de reemplazo
        """
        if not runs:
            return
        
        # Definir los marcadores y sus reemplazos
        replacements = {
            "{{fecha_evento}}": data.get("fecha_evento", ""),
            "{{hora_evento}}": data.get("hora_evento", ""),
            "{{lugar_evento}}": data.get("lugar_evento", ""),
            "{{referencia_evento}}": data.get("referencia_evento", "")
        }
        
        # Construir el texto completo para detectar marcadores
        full_text = ''.join(run.text for run in runs)
        
        # Verificar si hay marcadores
        has_markers = any(marker in full_text for marker in replacements.keys())
        
        if not has_markers:
            # No hay marcadores, no tocar nada
            return
        
        # Estrategia mejorada: Procesar run por run, pero mantener un buffer
        # para detectar marcadores que cruzan múltiples runs
        
        i = 0
        while i < len(runs):
            run = runs[i]
            
            # Si el run no tiene texto, verificar si tiene otros elementos importantes
            if not run.text:
                # Verificar si tiene saltos de línea u otros elementos
                if hasattr(run, '_element'):
                    # Buscar elementos especiales (breaks, imágenes, etc.)
                    has_special = False
                    if run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br') is not None:
                        has_special = True
                    if run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                        has_special = True
                    
                    if has_special:
                        # Preservar este run, no tocarlo
                        i += 1
                        continue
            
            # Construir texto desde este run hacia adelante
            text_buffer = ""
            run_group = []
            j = i
            
            while j < len(runs):
                current_run = runs[j]
                text_buffer += current_run.text
                run_group.append(j)
                
                # Verificar si hemos completado algún marcador
                found_complete_marker = False
                for marker in replacements.keys():
                    if marker in text_buffer:
                        found_complete_marker = True
                        break
                
                # Si encontramos un marcador completo o llegamos a un run vacío con elementos especiales
                if found_complete_marker:
                    break
                
                # Si el siguiente run está vacío y tiene elementos especiales, parar aquí
                if j + 1 < len(runs) and not runs[j + 1].text:
                    if hasattr(runs[j + 1], '_element'):
                        if runs[j + 1]._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br') is not None:
                            break
                
                j += 1
            
            # Hacer reemplazos en el buffer
            replaced_text = text_buffer
            for marker, value in replacements.items():
                replaced_text = replaced_text.replace(marker, value)
            
            # Si hubo cambios, actualizar los runs
            if replaced_text != text_buffer and run_group:
                # Poner el texto reemplazado en el primer run del grupo
                runs[run_group[0]].text = replaced_text
                # Limpiar los demás runs del grupo
                for idx in run_group[1:]:
                    runs[idx].text = ''
            
            # Avanzar al siguiente grupo
            i = j + 1
    
    def _replace_in_textboxes(self, doc, data: Dict[str, str]):
        """
        Reemplaza marcadores en cuadros de texto (textboxes).
        
        Los cuadros de texto son elementos especiales en Word que contienen
        párrafos pero no son accesibles directamente a través de doc.paragraphs.
        
        Args:
            doc: Documento de Word
            data: Diccionario con los valores de reemplazo
        """
        # Buscar todos los elementos txbxContent en el documento
        body = doc._element.body
        txbx_elements = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')
        
        for txbx in txbx_elements:
            # Buscar párrafos dentro del cuadro de texto
            paragraphs = txbx.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            
            for p_element in paragraphs:
                # Buscar runs dentro del párrafo
                run_elements = p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                
                # Crear una lista de objetos run simulados para usar con _replace_in_runs
                class TextBoxRun:
                    def __init__(self, run_element):
                        self._element = run_element
                        # Buscar elementos de texto
                        text_elements = run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        self._text_elements = text_elements
                        self._text = ''.join(t.text or '' for t in text_elements)
                    
                    @property
                    def text(self):
                        return self._text
                    
                    @text.setter
                    def text(self, value):
                        self._text = value
                        # Actualizar el primer elemento de texto, eliminar los demás
                        if self._text_elements:
                            self._text_elements[0].text = value
                            # Eliminar elementos de texto adicionales
                            for t_elem in self._text_elements[1:]:
                                parent = t_elem.getparent()
                                if parent is not None:
                                    parent.remove(t_elem)
                        elif value:
                            # Si no hay elementos de texto, crear uno
                            from docx.oxml import parse_xml
                            t_elem = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{value}</w:t>')
                            self._element.append(t_elem)
                
                # Crear objetos run para este párrafo
                runs = [TextBoxRun(r) for r in run_elements]
                
                # Aplicar reemplazos
                if runs:
                    self._replace_in_runs(runs, data)
