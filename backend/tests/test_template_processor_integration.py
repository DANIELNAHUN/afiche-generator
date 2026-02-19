"""
Integration tests for TemplateProcessor with real template files

Tests the template processor with the actual Formato a4.docx and Formato 4x1.docx files.
"""

import pytest
import os
import tempfile
from docx import Document
from services.template_processor import TemplateProcessor


class TestTemplateProcessorIntegration:
    """Integration tests with real template files"""
    
    def test_process_a4_template_with_real_file(self):
        """Verifica que puede procesar la plantilla A4 real"""
        processor = TemplateProcessor()
        
        # Verificar que la plantilla existe
        template_path = os.path.join(processor.template_dir, processor.templates["a4"])
        if not os.path.exists(template_path):
            pytest.skip(f"Plantilla A4 no encontrada en {template_path}")
        
        # Crear archivo de salida temporal
        with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name
        
        try:
            data = {
                "fecha_evento": "15 de Diciembre, 2024",
                "hora_evento": "7:00 PM",
                "lugar_evento": "Auditorio Central",
                "referencia_evento": "Campaña Navidad"
            }
            
            result_path = processor.process_template("a4", data, output_path)
            
            # Verificar que el archivo fue creado
            assert os.path.exists(result_path)
            
            # Verificar que es un documento Word válido
            doc = Document(output_path)
            assert doc is not None
            
            # Verificar que tiene contenido
            assert len(doc.paragraphs) > 0 or len(doc.tables) > 0
            
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
    
    def test_process_4x1_template_with_real_file(self):
        """Verifica que puede procesar la plantilla 4x1 real"""
        processor = TemplateProcessor()
        
        # Verificar que la plantilla existe
        template_path = os.path.join(processor.template_dir, processor.templates["4x1"])
        if not os.path.exists(template_path):
            pytest.skip(f"Plantilla 4x1 no encontrada en {template_path}")
        
        # Crear archivo de salida temporal
        with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name
        
        try:
            data = {
                "fecha_evento": "20 de Enero, 2025",
                "hora_evento": "6:00 PM",
                "lugar_evento": "Iglesia Central",
                "referencia_evento": ""  # Campo opcional vacío
            }
            
            result_path = processor.process_template("4x1", data, output_path)
            
            # Verificar que el archivo fue creado
            assert os.path.exists(result_path)
            
            # Verificar que es un documento Word válido
            doc = Document(output_path)
            assert doc is not None
            
            # Verificar que tiene contenido
            assert len(doc.paragraphs) > 0 or len(doc.tables) > 0
            
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
    
    def test_process_template_preserves_formatting(self):
        """Verifica que el procesamiento preserva el formato original"""
        processor = TemplateProcessor()
        
        template_path = os.path.join(processor.template_dir, processor.templates["a4"])
        if not os.path.exists(template_path):
            pytest.skip(f"Plantilla A4 no encontrada en {template_path}")
        
        # Leer la plantilla original
        original_doc = Document(template_path)
        original_paragraph_count = len(original_doc.paragraphs)
        original_table_count = len(original_doc.tables)
        
        # Procesar la plantilla
        with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_output:
            output_path = temp_output.name
        
        try:
            data = {
                "fecha_evento": "Test Date",
                "hora_evento": "Test Time",
                "lugar_evento": "Test Location",
                "referencia_evento": "Test Reference"
            }
            
            processor.process_template("a4", data, output_path)
            
            # Leer el documento procesado
            processed_doc = Document(output_path)
            
            # Verificar que la estructura se preservó
            assert len(processed_doc.paragraphs) == original_paragraph_count
            assert len(processed_doc.tables) == original_table_count
            
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
