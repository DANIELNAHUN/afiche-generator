"""
Unit tests for TemplateProcessor

Tests the template processing functionality including:
- Reading Word templates
- Replacing markers in paragraphs and tables
- Handling optional fields
- Error handling for missing templates
"""

import pytest
import os
import tempfile
from docx import Document
from services.template_processor import TemplateProcessor


class TestTemplateProcessor:
    """Test suite for TemplateProcessor class"""
    
    def test_init_sets_template_configuration(self):
        """Verifica que la inicialización configura las plantillas correctamente"""
        processor = TemplateProcessor()
        
        assert processor.template_dir is not None
        assert "a4" in processor.templates
        assert "4x1" in processor.templates
        assert processor.templates["a4"] == "Formato a4.docx"
        assert processor.templates["4x1"] == "Formato 4x1.docx"
    
    def test_process_template_with_invalid_type_raises_error(self):
        """Verifica que un tipo de plantilla inválido genera KeyError"""
        processor = TemplateProcessor()
        
        with pytest.raises(KeyError, match="Tipo de plantilla inválido"):
            processor.process_template(
                "invalid_type",
                {"fecha_evento": "test"},
                "output.docx"
            )
    
    def test_process_template_with_missing_file_raises_error(self):
        """Verifica que una plantilla inexistente genera FileNotFoundError"""
        processor = TemplateProcessor()
        # Modificar temporalmente la ruta de la plantilla
        processor.templates["a4"] = "nonexistent_template.docx"
        
        with pytest.raises(FileNotFoundError, match="Plantilla no encontrada"):
            processor.process_template(
                "a4",
                {"fecha_evento": "test"},
                "output.docx"
            )
    
    def test_replace_in_runs_replaces_all_markers(self):
        """Verifica que _replace_in_runs reemplaza todos los marcadores"""
        processor = TemplateProcessor()
        
        # Crear un documento de prueba en memoria
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Fecha: {{fecha_evento}}, Hora: {{hora_evento}}")
        
        data = {
            "fecha_evento": "15 de Diciembre",
            "hora_evento": "7:00 PM",
            "lugar_evento": "Auditorio Central",
            "referencia_evento": "Ref123"
        }
        
        processor._replace_in_runs(paragraph.runs, data)
        
        # Verificar que los marcadores fueron reemplazados
        text = "".join(run.text for run in paragraph.runs)
        assert "{{fecha_evento}}" not in text
        assert "{{hora_evento}}" not in text
        assert "15 de Diciembre" in text
        assert "7:00 PM" in text
    
    def test_replace_in_runs_handles_empty_optional_field(self):
        """Verifica que _replace_in_runs maneja campos opcionales vacíos"""
        processor = TemplateProcessor()
        
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Referencia: {{referencia_evento}}")
        
        data = {
            "fecha_evento": "15 de Diciembre",
            "hora_evento": "7:00 PM",
            "lugar_evento": "Auditorio Central",
            "referencia_evento": ""  # Campo opcional vacío
        }
        
        processor._replace_in_runs(paragraph.runs, data)
        
        text = "".join(run.text for run in paragraph.runs)
        assert "{{referencia_evento}}" not in text
        assert text == "Referencia: "
    
    def test_replace_in_runs_handles_missing_keys(self):
        """Verifica que _replace_in_runs maneja claves faltantes con valores vacíos"""
        processor = TemplateProcessor()
        
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Lugar: {{lugar_evento}}")
        
        # Data sin la clave lugar_evento
        data = {
            "fecha_evento": "15 de Diciembre",
            "hora_evento": "7:00 PM"
        }
        
        processor._replace_in_runs(paragraph.runs, data)
        
        text = "".join(run.text for run in paragraph.runs)
        # Debe reemplazar con cadena vacía
        assert "{{lugar_evento}}" not in text
        assert text == "Lugar: "
    
    def test_process_template_creates_output_file(self):
        """Verifica que process_template crea el archivo de salida"""
        processor = TemplateProcessor()
        
        # Crear una plantilla de prueba temporal
        with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_template:
            temp_template_path = temp_template.name
        
        try:
            # Crear un documento Word simple
            doc = Document()
            doc.add_paragraph("Evento: {{fecha_evento}} a las {{hora_evento}}")
            doc.save(temp_template_path)
            
            # Configurar el procesador para usar la plantilla temporal
            processor.templates["a4"] = os.path.basename(temp_template_path)
            processor.template_dir = os.path.dirname(temp_template_path)
            
            # Crear archivo de salida temporal
            with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_output:
                output_path = temp_output.name
            
            try:
                data = {
                    "fecha_evento": "20 de Enero",
                    "hora_evento": "6:00 PM",
                    "lugar_evento": "Iglesia Central",
                    "referencia_evento": ""
                }
                
                result_path = processor.process_template("a4", data, output_path)
                
                # Verificar que el archivo fue creado
                assert os.path.exists(result_path)
                assert result_path == output_path
                
                # Verificar que el contenido fue procesado
                processed_doc = Document(output_path)
                text = "\n".join(p.text for p in processed_doc.paragraphs)
                assert "{{fecha_evento}}" not in text
                assert "{{hora_evento}}" not in text
                assert "20 de Enero" in text
                assert "6:00 PM" in text
                
            finally:
                # Limpiar archivo de salida
                if os.path.exists(output_path):
                    os.remove(output_path)
        
        finally:
            # Limpiar plantilla temporal
            if os.path.exists(temp_template_path):
                os.remove(temp_template_path)
    
    def test_process_template_handles_tables(self):
        """Verifica que process_template reemplaza marcadores en tablas"""
        processor = TemplateProcessor()
        
        # Crear una plantilla con tabla
        with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_template:
            temp_template_path = temp_template.name
        
        try:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Fecha"
            table.cell(0, 1).text = "{{fecha_evento}}"
            table.cell(1, 0).text = "Lugar"
            table.cell(1, 1).text = "{{lugar_evento}}"
            doc.save(temp_template_path)
            
            processor.templates["a4"] = os.path.basename(temp_template_path)
            processor.template_dir = os.path.dirname(temp_template_path)
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as temp_output:
                output_path = temp_output.name
            
            try:
                data = {
                    "fecha_evento": "25 de Marzo",
                    "hora_evento": "5:00 PM",
                    "lugar_evento": "Auditorio Principal",
                    "referencia_evento": ""
                }
                
                processor.process_template("a4", data, output_path)
                
                # Verificar que los marcadores en la tabla fueron reemplazados
                processed_doc = Document(output_path)
                table = processed_doc.tables[0]
                
                cell_text = table.cell(0, 1).text
                assert "{{fecha_evento}}" not in cell_text
                assert "25 de Marzo" in cell_text
                
                cell_text = table.cell(1, 1).text
                assert "{{lugar_evento}}" not in cell_text
                assert "Auditorio Principal" in cell_text
                
            finally:
                if os.path.exists(output_path):
                    os.remove(output_path)
        
        finally:
            if os.path.exists(temp_template_path):
                os.remove(temp_template_path)
